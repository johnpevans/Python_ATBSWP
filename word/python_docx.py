#! python3

#From the exercise in Automate the Boring Stuff with Python

import docx

insert_path = "C:\\[INSERT_PATH]"

d = docx.Document('{}\\documents\\demo.docx'.format(insert_path))

p = d.paragraphs[1]

p.runs[1].text

p.runs[1].italic=True
p.runs[0].underline=True

d.save('{}\\documents\\demo2.docx'.format(insert_path))

p.style = 'Title'

d.save('insert_path\\documents\\demo3.docx'.format(insert_path))
#New word Document
d= docx.Document()
d.add_paragraph('Hello this is a paragraph.')
d.add_paragraph('Hello this is another paragraph.')
d.save('{}\\documents\\demo4.docx'.format(insert_path))

p0 = d.paragraphs[0]
p0.add_run(' This is a new run.')
p0.runs[1].bold = True
d.save('{}\\documents\\demo5.docx'.format(insert_path))

#Get text from document as a string

def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

print(getText('{}\\documents\\demo.docx'.format(insert_path)))
